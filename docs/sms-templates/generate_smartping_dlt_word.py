"""
Generate SmartPing (smartping.live) DLT registration Word pack for Starr Kuts / BillVyapp.

Variable format: {#VariableName#}
Header (Sender ID): STRKUT

CRITICAL (STPL rejection 03-Aug-2026):
  "Entity brand name is not mentioned in the SMS content…"
  → Brand MUST be hardcoded literal text in every template body.
  → Do NOT use {#BusinessName#} for the entity brand.
  → Fill "Brand Name" on the SmartPing form (must match entity registration).

Run: python docs/sms-templates/generate_smartping_dlt_word.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parent
OUT = ROOT / "BillVyapp-SmartPing-DLT-Templates.docx"
HEADER = "STRKUT"
# Must match the Principal Entity brand registered on SmartPing / DLT.
# If your portal brand differs (e.g. "StarrKuts" / "STRKUT"), change BRAND to that exact string.
BRAND = "Starr Kuts"

TEMPLATES = [
    {
        "id": "SMS-01",
        "name": "Login / Registration OTP",
        "template_type": "Transactional",
        "category": "Authentication",
        "priority": "P0",
        "where": "Auth login OTP + registration verification",
        "body": (
            f"{BRAND}: Your verification code is {{#OTP#}}. "
            "Valid for {#ExpiryMinutes#} minutes. Do not share this code."
        ),
        "variables": [
            ("OTP", "482917"),
            ("ExpiryMinutes", "10"),
        ],
        "notes": "Brand is fixed prefix. OTP vars only.",
    },
    {
        "id": "SMS-02",
        "name": "Appointment Confirmation",
        "template_type": "Service Implicit",
        "category": "Appointments",
        "priority": "P0",
        "where": "Appointments confirm / notify",
        "body": (
            f"Dear Customer, your appointment with {BRAND} is confirmed for "
            "{#Date#} at {#Time#}. For assistance, contact {#Phone#}. Thank you. - "
            f"{BRAND}"
        ),
        "variables": [
            ("Date", "03-Aug-2026"),
            ("Time", "11:30 AM"),
            ("Phone", "9876543210"),
        ],
        "notes": "FIXED brand (not a variable). This is the template that was rejected when brand was {#var#}.",
    },
    {
        "id": "SMS-03",
        "name": "Appointment Reminder",
        "template_type": "Service Implicit",
        "category": "Appointments",
        "priority": "P0",
        "where": "Automated reminder job + Appointments notify",
        "body": (
            f"Dear Customer, this is a reminder that your appointment with {BRAND} "
            "is scheduled on {#Date#} at {#Time#}. We look forward to serving you. - "
            f"{BRAND}"
        ),
        "variables": [
            ("Date", "03-Aug-2026"),
            ("Time", "11:30 AM"),
        ],
        "notes": "FIXED brand.",
    },
    {
        "id": "SMS-04",
        "name": "Running Late / Delay",
        "template_type": "Service Implicit",
        "category": "Appointments",
        "priority": "P1",
        "where": "Appointments notify — delay",
        "body": (
            f"Dear Customer, {BRAND} regrets to inform you that your appointment is delayed by "
            "approximately {#DelayTime#}. We apologize for the inconvenience and appreciate "
            f"your patience. - {BRAND}"
        ),
        "variables": [
            ("DelayTime", "15 minutes"),
        ],
        "notes": "Brand added — previous draft had no brand (would also be rejected).",
    },
    {
        "id": "SMS-05",
        "name": "Thank You After Visit",
        "template_type": "Promotional",
        "category": "Appointments",
        "priority": "P1",
        "where": "Post checkout / appointment completed",
        "body": (
            f"Dear Customer, thank you for visiting {BRAND} today. "
            f"We appreciate your trust and look forward to serving you again. - {BRAND}"
        ),
        "variables": [],
        "notes": "Promotional (SE retired). No variables — brand fixed.",
    },
    {
        "id": "SMS-06",
        "name": "Upcoming Appointment (Customers)",
        "template_type": "Service Implicit",
        "category": "Customers",
        "priority": "P1",
        "where": "Customers module notify",
        "body": (
            f"Dear Customer, your upcoming appointment with {BRAND} is on "
            "{#Date#} at {#Time#}. Please arrive {#Minutes#} minutes early. - "
            f"{BRAND}"
        ),
        "variables": [
            ("Date", "03-Aug-2026"),
            ("Time", "11:30 AM"),
            ("Minutes", "10"),
        ],
        "notes": "FIXED brand.",
    },
    {
        "id": "SMS-07",
        "name": "Loyalty Points Reminder",
        "template_type": "Promotional",
        "category": "Customers / Loyalty",
        "priority": "P1",
        "where": "Customers loyalty notify",
        "body": (
            "Dear Customer, you have {#Points#} loyalty points available in your account. "
            f"Redeem them before {{#ExpiryDate#}}. Thank you for choosing {BRAND}."
        ),
        "variables": [
            ("Points", "450"),
            ("ExpiryDate", "31-Dec-2026"),
        ],
        "notes": "FIXED brand at end.",
    },
    {
        "id": "SMS-08",
        "name": "Membership / Member Offer",
        "template_type": "Promotional",
        "category": "Customers / Loyalty",
        "priority": "P1",
        "where": "Customers Send Member Offer",
        "body": (
            "Dear Customer, enjoy an exclusive member offer of {#OfferDetails#} valid until "
            f"{{#ExpiryDate#}}. Visit {BRAND} to redeem your benefit. - {BRAND}"
        ),
        "variables": [
            ("OfferDetails", "20 percent OFF services"),
            ("ExpiryDate", "31-Aug-2026"),
        ],
        "notes": "FIXED brand. Prefer 'percent' over % if DLT is picky.",
    },
    {
        "id": "SMS-09",
        "name": "Birthday Wish + Offer",
        "template_type": "Promotional",
        "category": "Customers",
        "priority": "P1",
        "where": "Birthday job / Customers birthday",
        "body": (
            "Happy Birthday, {#CustomerName#}! Wishing you a wonderful year ahead. "
            "Enjoy our special offer: {#OfferDetails#}, valid until {#ExpiryDate#}. "
            f"Team {BRAND}."
        ),
        "variables": [
            ("CustomerName", "Priya Sharma"),
            ("OfferDetails", "20 percent OFF"),
            ("ExpiryDate", "31-Aug-2026"),
        ],
        "notes": "FIXED brand via Team Starr Kuts.",
    },
    {
        "id": "SMS-10",
        "name": "Payment Received Receipt",
        "template_type": "Transactional",
        "category": "Billing / Receipts",
        "priority": "P0",
        "where": "Billing checkout / walk-in receipt",
        "body": (
            f"Dear Customer, {BRAND} has received your payment of Rs.{{#Amount#}} against Invoice "
            "{#InvoiceNo#} on {#Date#}. Thank you for your payment. - "
            f"{BRAND}"
        ),
        "variables": [
            ("Amount", "1850"),
            ("InvoiceNo", "INV-2026-0842"),
            ("Date", "03-Aug-2026"),
        ],
        "notes": "Brand hardcoded — payment templates previously had no brand.",
    },
    {
        "id": "SMS-11",
        "name": "Payment Reminder (Due)",
        "template_type": "Transactional",
        "category": "Finance",
        "priority": "P1",
        "where": "Finance pending payments",
        "body": (
            f"Dear Customer, a payment of Rs.{{#Amount#}} for Invoice {{#InvoiceNo#}} at {BRAND} "
            "is due on {#DueDate#}. Kindly make the payment to avoid any inconvenience. - "
            f"{BRAND}"
        ),
        "variables": [
            ("Amount", "1850"),
            ("InvoiceNo", "INV-2026-0842"),
            ("DueDate", "20-Jul-2026"),
        ],
        "notes": "FIXED brand.",
    },
    {
        "id": "SMS-12",
        "name": "Payment Reminder (Overdue)",
        "template_type": "Transactional",
        "category": "Finance",
        "priority": "P1",
        "where": "Finance overdue payments",
        "body": (
            f"Dear Customer, your payment of Rs.{{#Amount#}} for Invoice {{#InvoiceNo#}} at {BRAND} "
            "is overdue since {#DueDate#}. Kindly clear the outstanding amount at the earliest. "
            f"Thank you. - {BRAND}"
        ),
        "variables": [
            ("Amount", "1850"),
            ("InvoiceNo", "INV-2026-0842"),
            ("DueDate", "20-Jul-2026"),
        ],
        "notes": "FIXED brand.",
    },
    {
        "id": "SMS-13",
        "name": "Coupon Send (Single / Bulk)",
        "template_type": "Promotional",
        "category": "Coupons / Promotions",
        "priority": "P1",
        "where": "Coupons + Customers coupon send",
        "body": (
            "Dear Customer, your coupon code {#CouponCode#} is worth {#OfferDetails#} and is "
            f"valid until {{#ExpiryDate#}}. Redeem it at {BRAND}."
        ),
        "variables": [
            ("CouponCode", "SAVE20"),
            ("OfferDetails", "20 percent OFF"),
            ("ExpiryDate", "31-Aug-2026"),
        ],
        "notes": "FIXED brand.",
    },
    {
        "id": "SMS-14",
        "name": "Request Feedback / Review",
        "template_type": "Promotional",
        "category": "Feedback",
        "priority": "P2",
        "where": "Feedback request modal",
        "body": (
            f"Dear Customer, thank you for choosing {BRAND}. Please share your "
            "feedback here: {#FeedbackLink#}. Your opinion helps us improve. - "
            f"{BRAND}"
        ),
        "variables": [
            ("FeedbackLink", "https://billvy.app/r/abc123"),
        ],
        "notes": "Promotional (SE retired). FIXED brand.",
    },
    {
        "id": "SMS-15",
        "name": "New Appointment Booking Notification",
        "template_type": "Transactional",
        "category": "Appointments (Staff)",
        "priority": "P0",
        "where": "On new appointment create — staff / salon phone",
        "body": (
            f"Dear {{#StaffName#}}, a new appointment at {BRAND} has been booked by "
            "{#CustomerName#} for {#Date#} at {#Time#}. Please review your schedule. - "
            f"{BRAND}"
        ),
        "variables": [
            ("StaffName", "Ananya"),
            ("CustomerName", "Priya Sharma"),
            ("Date", "03-Aug-2026"),
            ("Time", "11:30 AM"),
        ],
        "notes": "Staff-facing. FIXED brand.",
    },
]


def set_run_font(run, size=11, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = "Calibri"
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), "Calibri")
    rFonts.set(qn("w:hAnsi"), "Calibri")
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading_line(doc: Document, text: str, size=16):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=True, color=(17, 17, 24))


def add_label_value(doc: Document, label: str, value: str):
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}: ")
    set_run_font(r1, size=11, bold=True)
    r2 = p.add_run(value)
    set_run_font(r2, size=11)


def fill_sample(body: str, variables: list[tuple[str, str]]) -> str:
    out = body
    for name, sample in variables:
        out = out.replace("{#" + name + "#}", sample)
    return out


def build() -> Path:
    doc = Document()

    add_heading_line(doc, "BillVyapp / Starr Kuts — SmartPing DLT SMS Templates (v2)", 18)
    p = doc.add_paragraph()
    set_run_font(
        p.add_run(
            f"Provider: SmartPing · Header: {HEADER} · Entity brand (FIXED in every SMS): {BRAND}"
        ),
        size=10,
        color=(100, 100, 100),
    )

    warn = doc.add_paragraph()
    set_run_font(
        warn.add_run(
            "WHY SMS-02 WAS REJECTED (STPL, Ref 11-D1K6MSD7F7SN): "
            '"Entity brand name is not mentioned in the SMS content…". '
            "Putting the brand only in the SAMPLE (via {#BusinessName#}) does NOT count. "
            "Brand must appear as fixed text in the template body. "
            "Also fill Brand Name on the SmartPing form (yours showed \"-\"). "
            "Repeated wrong submissions can blacklist header STRKUT."
        ),
        size=10,
        bold=True,
        color=(160, 40, 40),
    )

    checklist = [
        f"Brand Name field on form = exactly \"{BRAND}\" (or your DLT-registered entity name)",
        f"Header Associated = {HEADER}",
        "Every template body must contain the brand as plain text (not a variable)",
        "Only customer/date/amount/etc. fields are {#var#}",
        "Thank You + Feedback = Promotional (Service Explicit retired)",
        "Do not resubmit the old bodies that use {#BusinessName#} for the brand",
    ]
    add_heading_line(doc, "Before you resubmit", 14)
    for item in checklist:
        cp = doc.add_paragraph(style="List Number")
        set_run_font(cp.add_run(item), size=11)

    add_label_value(doc, "Total templates", str(len(TEMPLATES)))
    doc.add_paragraph()

    add_heading_line(doc, "Quick index", 14)
    for t in TEMPLATES:
        line = doc.add_paragraph(style="List Number")
        set_run_font(
            line.add_run(f"{t['id']} — {t['name']} [{t['template_type']}] ({t['priority']})"),
            size=10,
        )

    doc.add_page_break()

    for t in TEMPLATES:
        add_heading_line(doc, f"{t['id']}: {t['name']}", 14)
        add_label_value(doc, "Template Type (SmartPing)", t["template_type"])
        add_label_value(doc, "Category", t["category"])
        add_label_value(doc, "Priority", t["priority"])
        add_label_value(doc, "Where in BillVyapp", t["where"])
        add_label_value(doc, "Header Associated", HEADER)
        add_label_value(doc, "Brand Name (form field)", BRAND)
        if t["notes"]:
            add_label_value(doc, "Notes", t["notes"])

        body_label = doc.add_paragraph()
        set_run_font(body_label.add_run("Template text (copy into SmartPing):"), size=11, bold=True)
        body_p = doc.add_paragraph()
        set_run_font(body_p.add_run(t["body"]), size=11)

        vars_label = doc.add_paragraph()
        set_run_font(vars_label.add_run("Variables + sample values:"), size=11, bold=True)
        if not t["variables"]:
            vp = doc.add_paragraph(style="List Bullet")
            set_run_font(vp.add_run("None (fixed text only)"), size=10)
        for name, sample in t["variables"]:
            vp = doc.add_paragraph(style="List Bullet")
            set_run_font(vp.add_run(f"{{#{name}#}} → sample: {sample}"), size=10)

        sample_label = doc.add_paragraph()
        set_run_font(sample_label.add_run("Filled sample:"), size=11, bold=True)
        sample_p = doc.add_paragraph()
        set_run_font(
            sample_p.add_run(fill_sample(t["body"], t["variables"])),
            size=10,
            color=(60, 60, 60),
        )
        doc.add_paragraph()

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    path = build()
    print(f"Wrote {path}")
